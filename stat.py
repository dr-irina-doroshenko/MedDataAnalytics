import pandas as pd
import numpy as np
from itertools import combinations
from scipy import stats
from scipy.stats import mannwhitneyu, chi2_contingency, fisher_exact, kruskal, spearmanr, kendalltau
from docx import Document
from pathlib import Path

#загружам базу данных из Листа 1 Excel
INPUT_PATH = Path("data") / "Headache_Sample_DB.xlsx"
OUTPUT_PATH = Path("output") / "статистический_отчет.docx"
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

df = pd.read_excel(INPUT_PATH, sheet_name="Лист1")
#если с компьютера
#df = pd.read_excel(r"D:\Project\Headache_Sample_DB.xlsx", sheet_name="Лист1")

# 1. ПРЕДОБРАБОТКА ДАННЫХ
# 1.1. фильтрация - в исследование включаем только детей в возрасте 12 лет и старше (например)
# df = df[df['Age'] >= 12]

# 1.2А сохраняем данные (до заполнения пропусков, если потребуется)
df_raw = df.copy()
# для функции расчета категориальных данных по абс. (n=)
# - словарь из имени столбца и количества заполенных строк (для n=)
VALID_N = {col: df_raw[col].notna().sum() for col in df_raw.columns}
# - cловарь по индексу столбца (
VALID_N_BY_IDX = {i: df_raw.iloc[:, i].notna().sum() for i in range(df_raw.shape[1])}

# # 1.2B заполняем пустые строки (при необходимости!!!)
# # иначе расчет пойдет по количетсву n= для столбца
# # если столбец с количественными данными - заполняем пустую строку медианой, если весь столбец пустой - 0
# for col in df.select_dtypes(include=["int64", "float64"]).columns:
#     if df[col].isna().sum() > 0:
#         if df[col].dropna().shape[0] > 0:
#             df[col] = df[col].fillna(df[col].median())
#         else:
#             df[col] = df[col].fillna(0)
# # если столбец с качественными данными - заполняем пустую строку "нет данных"
# for col in df.select_dtypes(include=["object", "category"]).columns:
#     df[col] = df[col].fillna("нет данных")
# # проверяем остались ли пропуски
# print("Осталось NaN:", df.isna().sum().sum())

# 1.3 создание новой переменной (при необходимости) - например - возрастной период
# def create_age_groups(df, age_col='Age'):
#     """
#     Создаем переменную 'Возрастная_группа' на основе возраста
#     Периоды:
#     - < 1: грудной возраст
#     - 1-3: ранний возраст
#     - 3-5: дошкольный возраст
#     - 6-9: младший школьный возраст
#     - 10-134: средний школьный возраст
#     - >= 14: старший школьный возраст
#     """

#     def categorize_age(age):
#         if pd.isna(age):
#             return "нет данных"
#         elif age < 1:
#             return "грудной возраст"
#         elif 1 <= age < 3:
#             return "ранний возраст"
#         elif 3 <= age < 6:
#             return "дошкольный возраст"
#         elif 6 <= age < 10:
#             return "младший школьный возраст"
#         elif 10 <= age < 14:
#             return "средний школьный возраст"
#         else:  # age >= 14
#             return "старший школьный возраст"
    
#     # создаем новую колонку
#     df['Age_group'] = df[age_col].apply(categorize_age)
    
#     # Проверяем распределение
#     print("Распределение по возрастным группам:")
#     print(df['Age_group'].value_counts().sort_index())
    
#     return df
# # создаем возрастные группы в базе
# df = create_age_groups(df, age_col='Age')

#2. ВЫБОР ПЕРЕМЕННЫХ ДЛЯ АНАЛИЗА
# 2.1. выбор целевых переменных по имени столбца

# целевые переменные в данном файле - выделенные переменные для статистического анализа 
# !!!выбирает исследователь в ручную!
TARGET_COLUMNS = ['Gender', 'Age_group']

# 2.2. выбор столбцов для анализа

# определяем реальное количество столбцов в файле
total_cols = df.shape[1]
print(f"В файле {total_cols} столбцов")

# если нужно выбрать все
# cols_to_use0 = list(range(0, total_cols))

# выбираем в ручную, если база данных не адаптированы! 
# например без первого столбца ID[0] пациента (1-105)
cols_to_use0 = list(range(1, 105)) 
df_selected0 = df.iloc[:, cols_to_use0]

cols_to_use = list(range(0, 104)) 
df_base = df_selected0.iloc[:, cols_to_use].copy()  # базовый DataFrame без целевых переменных
df_base['Age_group'] = df['Age_group'] # базовый DataFrame с целевыми, которые мы создал cами

# 3. предобработка категориальных переменных
# определяем В РУЧНУЮ! столбы с категориальными переменными, если база не адаптирована!
# (например, ранговые переменные идут цифрами, а не текстом!!! - только вручную) 
# приводим их строкам (делаем их текстовым форматом)
categorical_idx = list(range(8, 104))
# categorical_idx = (
#     list(range(8, 195)) + 
#     list(range(324, 340, 2)) + 
#     list(range(352, 386, 2))
# )

for col in df_base.columns[categorical_idx]:
    df_base[col] = df_base[col].fillna("нет данных")
    df_base[col] = df_base[col].astype(str)


#3. CТАТАНАЛИЗ

# функция определения n= для категориальных столбцов
def get_expected_n(col_idx, default_n=None):
    """
    Определяем n для категориального столбца по его индексу.
    """
   
    return VALID_N_BY_IDX.get(col_idx, default_n if default_n is not None else len(df_raw))

# функция - сила связи по Крамеру для хи-квадрата
def calculate_cramers_v(chi2, n, min_dim):
    if n == 0 or min_dim <= 1:
        return 0
    return np.sqrt(chi2 / (n * (min_dim - 1)))

# функция - Хи-квадрат с поправкой Фишера, если >20% ячеек с ожидаемой частотй < 5 или минимум < 1
def chi2_with_fisher_correction(table, fisher_threshold=5):
    n = table.sum().sum()
    chi2, p, dof, expected = chi2_contingency(table)

    expected_less_5 = (expected < 5).sum()
    total_cells = expected.size
    percent_less_5 = expected_less_5 / total_cells
    min_expected = expected.min()

    use_fisher = percent_less_5 > 0.2 or min_expected < 1

    if use_fisher and table.shape == (2, 2):
        oddsratio, p_fisher = fisher_exact(table)
        return {
            'test_name': 'Фишер (точный)',
            'statistic': oddsratio,
            'p_value': p_fisher,
            'df': None,
            'chi2': None,
            'cramers_v': None,
            'interpretation': None,
            'used_fisher': True,
            'expected_less_5': expected_less_5,
            'percent_less_5': percent_less_5,
            'min_expected': min_expected
        }

    min_dim = min(table.shape)
    cramers_v = calculate_cramers_v(chi2, n, min_dim)

    return {
        'test_name': 'Хи-квадрат',
        'statistic': chi2,
        'p_value': p,
        'df': dof,
        'chi2': chi2,
        'cramers_v': cramers_v,
        'used_fisher': False,
        'expected_less_5': expected_less_5,
        'percent_less_5': percent_less_5,
        'min_expected': min_expected
    }


# функция - корреляционный анализ по Спирмену и Кендалл по типу все со всем
#! Спирмен оцениваем только - если хотя бы 1 переменная количественная
#! Кендал оцениваем - для количевеных vs ранговых, ранговых vs ранговых для бинарных - Фишер (см. выше)
def calculate_overall_correlations(df_base, num_cols):
    """
    Расчет корреляции Спирмена и Тау-Кендалла по всей базе
    """
    print("Корреляционный анализ по всей базе")

    spearman_matrix, spearman_sig = calculate_correlations(df_base, num_cols, method='spearman')
    kendall_matrix, kendall_sig = calculate_correlations(df_base, num_cols, method='kendall')
    
    corr_results = {
        'spearman': {'matrix': spearman_matrix, 'significant': spearman_sig},
        'kendall': {'matrix': kendall_matrix, 'significant': kendall_sig}
    }

    return corr_results

# функция - проводим общую описательную статистику
def calculate_overall_descriptive(df_base, categorical_idx, target_columns=None):
    """
    Расчет описательной статистики для всей базы данных
    """
    print("Описательной статистики для всей базы данных")

    n_total = len(df_base)
    print(f"Общее количество наблюдений: {n_total}")

    num_cols = df_base.select_dtypes(include=["int64","float64"]).columns.tolist()
    base_cat_cols = df_base.columns[categorical_idx].tolist()

    all_cat_cols = base_cat_cols.copy()

    if target_columns:
        for col in target_columns:
            if col in df_base.columns and col not in all_cat_cols:
                all_cat_cols.append(col)
                print(f"Добавлена целевая переменная: {col}")

    results = {
        'n_total': n_total,
        'numeric': {},
        'categorical': {}
    }

    # Количественные переменные
    print(f"\nКоличественные переменные ({len(num_cols)}):")
    for col in num_cols:
        data_valid = df_base[col].dropna()
        n_valid = len(data_valid)

        if n_valid == 0:
            print(f"  {col}: нет заполненных данных")
            continue

        stats = df_base[col].describe(percentiles=[0.25, 0.5, 0.75]).round(2)

        results['numeric'][col] = {
            'n': n_valid,
            'missing': n_total - n_valid,
            'mean': stats['mean'],
            'std': stats['std'],
            'median': stats['50%'],
            'q25': stats['25%'],
            'q75': stats['75%'],
            'min': stats['min'],
            'max': stats['max']
        }

        print(f"  {col} (n={n_valid}): M={stats['mean']:.2f}±{stats['std']:.2f}, "
              f"Me={stats['50%']:.2f}, Q1-Q3={stats['25%']:.2f}-{stats['75%']:.2f}, "
              f"min-max={stats['min']:.2f}-{stats['max']:.2f}")

    # Категориальные переменные
    print(f"\nКатегориальные переменные ({len(all_cat_cols)}):")
    for col in all_cat_cols:
        col_idx = df_base.columns.get_loc(col)
        expected_n = get_expected_n(col_idx, n_total)
        valid_data = df_base[col][(df_base[col].notna()) & (df_base[col] != "нет данных")]
        n_valid = len(valid_data)
        n_missing = n_total - n_valid
        if n_valid == 0:
            continue
        dist = valid_data.value_counts(normalize=True)
        cat_stats = []

        print(f"\n  {col} (n={expected_n}, заполнено: {n_valid}, пропусков/нет данных: {n_missing}):")

        for cat, prop in dist.items():
            abs_count = valid_data.value_counts().get(cat, 0)

            if expected_n > 0 and abs_count > 0:
                z = 1.96
                p = abs_count / expected_n
                n = expected_n

                denominator = 1 + z**2 / n
                centre = (p + z**2 / (2*n)) / denominator
                half_width = z * np.sqrt((p*(1-p) + z**2/(4*n)) / n) / denominator

                ci_low = max(0, round((centre - half_width) * 100, 1))
                ci_high = min(100, round((centre + half_width) * 100, 1))

                if ci_low > ci_high or ci_low < 0 or ci_high > 100:
                    alpha = 0.05
                    ci_low = round(stats.beta.ppf(alpha/2, abs_count, expected_n - abs_count + 1) * 100, 1)
                    ci_high = round(stats.beta.ppf(1 - alpha/2, abs_count + 1, expected_n - abs_count) * 100, 1)
            else:
                ci_low = 0.0
                ci_high = 100.0 if abs_count > 0 else 0.0

            pct = round((abs_count / expected_n) * 100, 1)

            cat_stats.append({
                'category': cat,
                'n': abs_count,
                'percent': pct,
                'ci_low': ci_low,
                'ci_high': ci_high
            })
            print(f"    {cat}: n={abs_count} {pct}% (95%ДИ: {ci_low}-{ci_high})")

        results['categorical'][col] = {
            'n_total': expected_n,
            'n_valid': n_valid,
            'n_missing': n_missing,
            'categories': cat_stats
        }

    # Корреляция по всей базе
    overall_corr = calculate_overall_correlations(df_base, num_cols)
    results['correlations'] = overall_corr

    return results


# функция - описательная статистика целевых переменных
def calculate_group_descriptive(df_work, target_column, groups, num_cols, cat_cols, n_total):
    """
    Расчет описательной статистики по группам целевой переменной
    """
    print(f"Описательная статистика по целевым переменным: {target_column}")

    group_stats = {}

    for g in groups:
        df_g = df_work[df_work[target_column].astype(str) == str(g)]
        n_g = len(df_g)

        print(f"\nГруппа {g} (n={n_g}):")
        group_stats[str(g)] = {'numeric': {}, 'categorical': {}}

        # Количественные переменные
        for col in num_cols:
            data = df_g[col].dropna()
            n_valid = len(data)

            if n_valid > 0:
                mean_val = data.mean()
                std_val = data.std() if n_valid > 1 else 0.0
                median_val = data.median()
                q25 = data.quantile(0.25)
                q75 = data.quantile(0.75)
                min_val = data.min()
                max_val = data.max()

                group_stats[str(g)]['numeric'][col] = {
                    'n': n_valid,
                    'mean': mean_val,
                    'std': std_val,
                    'median': median_val,
                    'q25': q25,
                    'q75': q75,
                    'min': min_val,
                    'max': max_val
                }

                print(f"  {col}: M={mean_val:.2f}±{std_val:.2f}, Ме={median_val:.2f}, "
                      f"Q1-Q3={q25:.2f}-{q75:.2f}, min-max={min_val:.2f}-{max_val:.2f}")

        # Категориальные переменные
        for col in cat_cols:
            col_idx = df_work.columns.get_loc(col)
            expected_n = get_expected_n(col_idx, n_g)  # за 100% принимается число заполненных ячеек в столбце (по всей базе)
            # expected_n = n_g                         # за 100% принимается размер группы (например, только мужчины)
            # expected_n = n_total                     # за 100% принимается общее число пациентов в базе

            valid_data = df_g[col][(df_g[col].notna()) & (df_g[col] != "нет данных")]
            n_valid = len(valid_data)

            if n_valid > 0:
                dist = valid_data.value_counts(normalize=True)
                cat_stats = []

                print(f"  {col} (n={expected_n}, заполнено: {n_valid}):")

                for cat, prop in dist.items():
                    abs_count = valid_data.value_counts().get(cat, 0)
                    if expected_n > 0 and abs_count > 0:
                        z = 1.96
                        p = abs_count / expected_n
                        n = expected_n

                        denominator = 1 + z**2 / n
                        centre = (p + z**2 / (2*n)) / denominator
                        half_width = z * np.sqrt((p*(1-p) + z**2/(4*n)) / n) / denominator

                        ci_low = max(0, round((centre - half_width) * 100, 1))
                        ci_high = min(100, round((centre + half_width) * 100, 1))

                        if ci_low > ci_high or ci_low < 0 or ci_high > 100:
                            alpha = 0.05
                            ci_low = round(stats.beta.ppf(alpha/2, abs_count, expected_n - abs_count + 1) * 100, 1)
                            ci_high = round(stats.beta.ppf(1 - alpha/2, abs_count + 1, expected_n - abs_count) * 100, 1)
                    else:
                        ci_low = 0.0
                        ci_high = 100.0 if abs_count > 0 else 0.0

                    pct = round((abs_count / expected_n) * 100, 1)

                    cat_stats.append({
                        'category': cat,
                        'n': abs_count,
                        'percent': pct,
                        'ci_low': ci_low,
                        'ci_high': ci_high
                    })

                    print(f"    {cat}: n={abs_count} {pct}% (95%ДИ: {ci_low}-{ci_high})")

                group_stats[str(g)]['categorical'][col] = {
                    'n_valid': n_valid,
                    'n_expected': expected_n,
                    'categories': cat_stats
                }

    return group_stats

# функция корреляционного анализа (и выгрузка)
def calculate_correlations(df_work, num_cols, method='spearman'):
    """
    Расчитываем корреляционную матрицу для числовых переменных
    """
    if len(num_cols) < 2:
        return None, None

    df_numeric = df_work[num_cols].apply(pd.to_numeric, errors='coerce')
    n = len(num_cols)
    
    # создаем пустые шаблоны для матриц
    corr_matrix = pd.DataFrame(np.zeros((n, n)), index=num_cols, columns=num_cols)
    p_matrix = pd.DataFrame(np.ones((n, n)), index=num_cols, columns=num_cols)

    if method == 'spearman':
        for i, col1 in enumerate(num_cols):
            for j, col2 in enumerate(num_cols):
                if i == j:
                    corr_matrix.iloc[i, j] = 1.0
                    p_matrix.iloc[i, j] = 0.0
                elif i < j:
                    try:
                        rho, p = spearmanr(df_numeric[col1].dropna(), df_numeric[col2].dropna())
                        corr_matrix.iloc[i, j] = rho
                        corr_matrix.iloc[j, i] = rho
                        p_matrix.iloc[i, j] = p
                        p_matrix.iloc[j, i] = p
                    except:
                        pass
        method_name = "Спирмена"

    elif method == 'kendall':
        for i, col1 in enumerate(num_cols):
            for j, col2 in enumerate(num_cols):
                if i == j:
                    corr_matrix.iloc[i, j] = 1.0
                    p_matrix.iloc[i, j] = 0.0
                elif i < j:
                    try:
                        tau, p = kendalltau(df_numeric[col1].dropna(), 
                                           df_numeric[col2].dropna())
                        corr_matrix.iloc[i, j] = tau
                        corr_matrix.iloc[j, i] = tau
                        p_matrix.iloc[i, j] = p
                        p_matrix.iloc[j, i] = p
                    except:
                        pass
        method_name = "Кендалла (тау)"
    else:
        return None, None

    print(f"\nКорреляционная матрица ({method_name}):")
    print(corr_matrix.round(3))

    significant = []
    for i in range(len(num_cols)):
        for j in range(i+1, len(num_cols)):
            col1, col2 = num_cols[i], num_cols[j]
            r = corr_matrix.iloc[i, j]
            p = p_matrix.iloc[i, j]
            if p < 0.05:
                significant.append({
                    'var1': col1,
                    'var2': col2,
                    'r': round(r, 3),
                    'p': round(p, 4)
                })
                print(f"  {col1} × {col2}: r={r:.3f}, p={p:.4f} ✓")

    return corr_matrix.round(3), significant

# функция - анализа для одной целевой переменной
def analyze_single_target(df_base, target_column, categorical_idx):
    """
    Проводит полный анализ для одной целевой переменной
    """
    print(f"Анализ целевой переменной: {target_column}")

    if target_column not in df_base.columns:
        print(f"Переменная '{target_column}' не найдена! Пропускаем.")
        return None

    df_selected = df_base.copy()
    df_selected['group'] = df_base[target_column].astype(str)

    groups = sorted(df_base[target_column].unique())
    print(f"Группы: {groups}")

    exclude_cols = [target_column, 'group']
    num_cols = df_selected.select_dtypes(include=["int64","float64"]).columns
    num_cols = [c for c in num_cols if c not in exclude_cols]

    cat_cols = df_selected.select_dtypes(include=["object","category"]).columns
    cat_cols = [c for c in cat_cols if c not in exclude_cols]

    df_work = df_selected.copy()

    # 1. Описательная статистика
    group_descriptive = calculate_group_descriptive(
        df_work, target_column, groups, num_cols, cat_cols, len(df_base)
    )

    # 2. Корреляционный анализ
    print("Корреляционный анализ групп целевой пемеменной")

    correlation_results = {}

    for g in groups:
        df_g = df_work[df_work['group'] == str(g)]
        print(f"\nГруппа {g} (n={len(df_g)}):")

        corr_spear, sig_spear = calculate_correlations(df_g, num_cols, method='spearman')
        corr_kend, sig_kend = calculate_correlations(df_g, num_cols, method='kendall')

        correlation_results[str(g)] = {
            'spearman': {'matrix': corr_spear, 'significant': sig_spear},
            'kendall': {'matrix': corr_kend, 'significant': sig_kend}
        }

    # 3. Множественное сравнение
    print("Множественное сравнение групп целевой пемеменной")
    
    mult_results = {'numeric': [], 'categorical': []}

    for col in num_cols:
        try:
            group_data = [df_work[df_work['group'] == str(g)][col].dropna().values for g in groups]
            
            all_combined = np.concatenate(group_data)
            if len(group_data) > 1 and np.unique(all_combined).size > 1:
                stat, p = kruskal(*group_data)
                mult_results['numeric'].append({
                    'variable': col,
                    'h_statistic': stat,
                    'p_value': p,
                    'significant': p < 0.05
                })
        except Exception as e:
            print(f"Ошибка в {col}: {e}")

    for col in cat_cols:
        try:
            table = pd.crosstab(df_work['group'], df_work[col])
            test_result = chi2_with_fisher_correction(table)

            mult_results['categorical'].append({
                'variable': col,
                **test_result
            })
        except Exception as e:
            print(f"Ошибка в {col}: {e}")

    # 4. Попарные сравнения

    print("Попарное сравнение групп целевой пемеменной")
    pairwise_results = []
    pairs = list(combinations(groups, 2))

    for z1, z2 in pairs:
        df1 = df_work[df_work['group'] == str(z1)]
        df2 = df_work[df_work['group'] == str(z2)]
        df_pair = df_work[df_work['group'].isin([str(z1), str(z2)])]

        pair_results = {
            'group1': z1,
            'group2': z2,
            'n1': len(df1),
            'n2': len(df2),
            'numeric': [],
            'categorical': []
        }

        for col in num_cols:
            try:
                d1 = df1[col].dropna()
                d2 = df2[col].dropna()
                if len(d1) > 0 and len(d2) > 0:
                    stat, p = mannwhitneyu(d1, d2, alternative="two-sided")
                    pair_results['numeric'].append({
                        'variable': col,
                        'statistic': stat,
                        'p_value': p,
                        'significant': p < 0.05
                    })
            except Exception as e:
                pass

        for col in cat_cols:
            try:
                table = pd.crosstab(df_pair['group'], df_pair[col])
                test_result = chi2_with_fisher_correction(table)

                top_cat = df_pair[col].value_counts(normalize=True).idxmax()
                top_pct = round(df_pair[col].value_counts(normalize=True).max()*100, 1)

                pair_results['categorical'].append({
                    'variable': col,
                    'top_category': top_cat,
                    'top_percent': top_pct,
                    **test_result
                })
            except Exception as e:
                pass

        pairwise_results.append(pair_results)

    return {
        'target': target_column,
        'groups': groups,
        'df_selected': df_selected,
        'num_cols': num_cols,
        'cat_cols': cat_cols,
        'group_descriptive': group_descriptive,
        'mult_results': mult_results,
        'pairwise_results': pairwise_results,
        'correlation_results': correlation_results
    }

# функция генерации отчета
# def generate_combined_report(all_results, overall_stats, filename="статистический_отчет.docx"):
def generate_combined_report(all_results, overall_stats):
    """
    Создает отчет: анализ по всей базе статистика + анализ по всем целевым переменным
    """
    doc = Document()
    doc.add_heading('Комплексный статистический анализ', 0)

    # РАЗДЕЛ 1: ОБЩАЯ ОПИСАТЕЛЬНАЯ СТАТИСТИКА 
    doc.add_heading('1. Общая описательная статистика', level=1)
    doc.add_paragraph(f"Общее количество наблюдений: {overall_stats['n_total']}")

    # Числовые переменные
    if overall_stats['numeric']:
        doc.add_heading('Числовые переменные:', level=2)
        for col, stats in overall_stats['numeric'].items():
            doc.add_paragraph(
                f"{col}: M={stats['mean']:.2f}±{stats['std']:.2f}, "
                f"медиана={stats['median']:.2f}, "
                f"Q1-Q3={stats['q25']:.2f}-{stats['q75']:.2f}, "
                f"min-max={stats['min']:.2f}-{stats['max']:.2f}"
            )

    # Категориальные переменные
    if overall_stats['categorical']:
        doc.add_heading('Категориальные переменные:', level=2)
        for col, data in overall_stats['categorical'].items():
            marker = " [ЦЕЛЕВАЯ ПЕРЕМЕННАЯ]" if col in TARGET_COLUMNS else ""

            doc.add_paragraph(
                f"{col}{marker} (n={data['n_total']}, пропущено/нет данных: {data['n_missing']}):", 
                style='List Bullet'
            )

            for cat in data['categories']:
                doc.add_paragraph(
                    f"{cat['category']}: n={cat['n']}, {cat['percent']}%, "
                    f"95%ДИ: {cat['ci_low']}-{cat['ci_high']}",
                    style='List Continue'
                )

    # КОРРЕЛЯЦИИ ПО ВСЕЙ БАЗЕ 
    if overall_stats.get('correlations'):
        doc.add_heading('1.1. Корреляционный анализ по всей базе', level=2)

        # Спирмен
        spear = overall_stats['correlations']['spearman']
        if spear and spear['significant']:
            doc.add_heading('Значимые корреляции Спирмена (p<0.05):', level=3)
            for corr in spear['significant']:
                doc.add_paragraph(
                    f"{corr['var1']} × {corr['var2']}: "
                    f"r={corr['r']:.3f}, p={corr['p']:.4f}"
                )
        else:
            doc.add_paragraph('Нет значимых корреляций по Спирмену (p<0.05)')

        # Кендалл
        kend = overall_stats['correlations']['kendall']
        if kend and kend['significant']:
            doc.add_heading('Значимые корреляции Кендалла (тау, p<0.05):', level=3)
            for corr in kend['significant']:
                doc.add_paragraph(
                    f"{corr['var1']} × {corr['var2']}: "
                    f"τ={corr['r']:.3f}, p={corr['p']:.4f}"
                )
        else:
            doc.add_paragraph('Нет значимых корреляций по Кендаллу (p<0.05)')

    doc.add_page_break()

    # РАЗДЕЛ 2: АНАЛИЗ ПО ЦЕЛЕВЫМ ПЕРЕМЕННЫМ
    for result in all_results:
        if result is None:
            continue

        target = result['target']
        groups = result['groups']
        group_descriptive = result['group_descriptive']
        mult_results = result['mult_results']
        pairwise_results = result['pairwise_results']
        corr_results = result['correlation_results']

        doc.add_heading(f'2. Анализ по переменной: {target}', level=1)
        doc.add_paragraph(f'Группы: {", ".join(map(str, groups))} (всего {len(groups)})')

        # 2.0 ОПИСАТЕЛЬНАЯ СТАТИСТИКА ПО ГРУППАМ
        doc.add_heading('2.0. Описательная статистика по группам', level=2)

        for g in groups:
            g_str = str(g)
            if g_str not in group_descriptive:
                continue

            stats_g = group_descriptive[g_str]

            doc.add_heading(f'Группа {g}:', level=3)

            # Числовые признаки по группе
            if stats_g['numeric']:
                doc.add_heading('Числовые признаки:', level=4)
                for col, stats in stats_g['numeric'].items():
                    doc.add_paragraph(
                        f"{col}: M={stats['mean']:.2f}±{stats['std']:.2f}, "
                        f"медиана={stats['median']:.2f}, "
                        f"Q1-Q3={stats['q25']:.2f}-{stats['q75']:.2f}, "
                        f"min-max={stats['min']:.2f}-{stats['max']:.2f}"
                    )

            # Категориальные признаки по группе
            if stats_g['categorical']:
                doc.add_heading('Категориальные признаки:', level=4)
                for col, data in stats_g['categorical'].items():
                    doc.add_paragraph(f"{col} (n={data['n_expected']}, заполнено: {data['n_valid']}):", style='List Bullet')
                    for cat in data['categories']:
                        doc.add_paragraph(
                            f"{cat['category']}: n={cat['n']}, {cat['percent']}%, "
                            f"95%ДИ: {cat['ci_low']}-{cat['ci_high']}",
                            style='List Continue'
                        )

        # 2.1 КОРРЕЛЯЦИОННЫЙ АНАЛИЗ ПО ГРУППАМ
        doc.add_heading('2.1. Корреляционный анализ по группам', level=2)

        for g in groups:
            g_str = str(g)
            if g_str not in corr_results:
                continue

            doc.add_heading(f'Группа {g}:', level=3)

            spear = corr_results[g_str]['spearman']
            if spear['significant']:
                doc.add_heading('Значимые корреляции Спирмена:', level=4)
                for corr in spear['significant']:
                    doc.add_paragraph(
                        f"{corr['var1']} × {corr['var2']}: "
                        f"r={corr['r']:.3f}, p={corr['p']:.4f}"
                    )
            else:
                doc.add_paragraph('Нет значимых корреляций по Спирмену (p<0.05)')

            kend = corr_results[g_str]['kendall']
            if kend['significant']:
                doc.add_heading('Значимые корреляции Кендалла (тау):', level=4)
                for corr in kend['significant']:
                    doc.add_paragraph(
                        f"{corr['var1']} × {corr['var2']}: "
                        f"τ={corr['r']:.3f}, p={corr['p']:.4f}"
                    )
            else:
                doc.add_paragraph('Нет значимых корреляций по Кендаллу (p<0.05)')

        # 2.2 МНОЖЕСТВЕННОЕ СРАВНЕНИЕ
        doc.add_heading('2.2. Множественное сравнение всех групп', level=2)

        doc.add_heading('Числовые признаки (Краскал-Уоллис):', level=3)
        for res in mult_results['numeric']:
            if not res['significant']:
                continue
            doc.add_paragraph(f"{res['variable']}: H={res['h_statistic']:.2f}, "
                            f"p={res['p_value']:.4f}")

        doc.add_heading('Категориальные признаки:', level=3)
        for res in mult_results['categorical']:
            if res['p_value'] >= 0.05:
                continue

            if res['used_fisher']:
                doc.add_paragraph(
                    f"{res['variable']}: точный тест Фишера, "
                    f"p={res['p_value']:.4f}"
                )
            else:
                doc.add_paragraph(
                    f"{res['variable']}: χ²={res['chi2']:.2f}, "
                    f"df={res['df']}, p={res['p_value']:.4f}, "
                    f"Cramér's V={res['cramers_v']:.3f}"
                )

        # 2.3 Попарные сравнения 
        doc.add_heading('2.3. Попарные сравнения', level=2)

        for pair in pairwise_results:
            g1, g2 = pair['group1'], pair['group2']
            doc.add_heading(f'{g1} vs {g2} (n₁={pair["n1"]}, n₂={pair["n2"]})', level=3)

            sig_num = [r for r in pair['numeric'] if r['significant']]
            if sig_num:
                doc.add_paragraph('Числовые (Манн-Уитни):', style='List Bullet')
                for r in sig_num:
                    doc.add_paragraph(f'{r["variable"]}: U={r["statistic"]:.2f}, p={r["p_value"]:.4f}',
                                    style='List Continue')

            sig_cat = [r for r in pair['categorical'] if r['p_value'] < 0.05]
            if sig_cat:
                doc.add_paragraph('Категориальные:', style='List Bullet')
                for r in sig_cat:
                    interpretation = r.get('interpretation', 'не определена')
                    cramers_v = r.get('cramers_v', None)

                    if r['used_fisher']:
                        doc.add_paragraph(
                            f'{r["variable"]}: точный тест Фишера, '
                            f'p={r["p_value"]:.4f} (топ: {r["top_category"]}, {r["top_percent"]}%)',
                            style='List Continue'
                        )
                    else:
                        if cramers_v is not None:
                            doc.add_paragraph(
                                f'{r["variable"]}: χ²={r["chi2"]:.2f}, '
                                f'p={r["p_value"]:.4f}, Cramér\'s V={cramers_v:.3f} '
                                f'({interpretation}) (топ: {r["top_category"]}, {r["top_percent"]}%)',
                                style='List Continue'
                            )
                        else:
                            doc.add_paragraph(
                                f'{r["variable"]}: χ²={r["chi2"]:.2f}, '
                                f'p={r["p_value"]:.4f} (топ: {r["top_category"]}, {r["top_percent"]}%)',
                                style='List Continue'
                            )
            else:
                doc.add_paragraph('Нет значимых различий', style='List Bullet')

        doc.add_page_break()

    doc.save(OUTPUT_PATH)
    print(f"Статистический отчет сохранен: {OUTPUT_PATH.resolve()}")
    return OUTPUT_PATH

# ОСНОВНОЙ ЗАПУСК

if __name__ == "__main__":

    # 1. Общая описательная статистика (включая корреляции по всей базе)
    overall_stats = calculate_overall_descriptive(
        df_base, 
        categorical_idx, 
        target_columns=TARGET_COLUMNS
    )
    # 2. Проводим комплексный анализ для всех целевых переменных
    all_results = []

    for target in TARGET_COLUMNS:
        result = analyze_single_target(df_base, target, categorical_idx)
        if result:
            all_results.append(result)

    # 3. Генерируем единый в формат Word
    if all_results:
        generate_combined_report(
            all_results=all_results,
            overall_stats=overall_stats,
        )
    else:
        print("Нет результатов для формирования отчета!")